"""Format training-plan-300h.docx: borders, typography, table polish."""

from __future__ import annotations

import shutil
import tempfile
from pathlib import Path

import pypandoc
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "training-plan-300h.md"
DOCX_PATH = ROOT / "training-plan-300h.docx"

FONT_NAME = "Calibri"
BODY_SIZE = Pt(11)
HEADING_COLOR = RGBColor(0x1F, 0x38, 0x64)
HEADER_FILL = "E8EEF4"
BORDER = {"sz": 4, "val": "single", "color": "7F7F7F", "space": "0"}


def set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        element = OxmlElement(f"w:{edge}")
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))
        tc_borders.append(element)
    tc_pr.append(tc_borders)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top=BORDER,
                bottom=BORDER,
                start=BORDER,
                end=BORDER,
            )


def set_run_font(run, *, bold: bool | None = None, size=None, color=None) -> None:
    run.font.name = FONT_NAME
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def format_paragraph(paragraph, *, size=None, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, bold=bold, size=size or BODY_SIZE)


def format_document_styles(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2)

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading"):
            for run in paragraph.runs:
                set_run_font(run, bold=True, color=HEADING_COLOR)
            paragraph.paragraph_format.space_before = Pt(14)
            paragraph.paragraph_format.space_after = Pt(6)
        elif style_name == "Quote":
            paragraph.paragraph_format.left_indent = Cm(0.8)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(4)
            for run in paragraph.runs:
                set_run_font(run, size=Pt(10), color=RGBColor(0x55, 0x55, 0x55))
        else:
            format_paragraph(paragraph)
            paragraph.paragraph_format.space_after = Pt(4)


def is_total_row(row) -> bool:
    text = " ".join(cell.text.strip() for cell in row.cells)
    return text.startswith("Итого") or " **Итого**" in text or text == "Итого"


def format_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")
    tbl_pr.append(tbl_w)

    set_table_borders(table)

    for row_idx, row in enumerate(table.rows):
        total_row = is_total_row(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                if row_idx == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        bold=row_idx == 0 or total_row or run.bold,
                        size=Pt(10) if len(table.columns) > 6 else BODY_SIZE,
                    )
            if row_idx == 0:
                shade_cell(cell, HEADER_FILL)
            elif total_row:
                shade_cell(cell, "F5F5F5")


def build_docx() -> None:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)

    try:
        pypandoc.convert_file(
            str(MD_PATH),
            "docx",
            outputfile=str(tmp_path),
            extra_args=["--standalone"],
        )
        doc = Document(str(tmp_path))
        format_document_styles(doc)
        for table in doc.tables:
            format_table(table)
        doc.save(str(tmp_path))
        try:
            shutil.copy2(tmp_path, DOCX_PATH)
        except PermissionError:
            fallback = DOCX_PATH.with_stem(DOCX_PATH.stem + "-formatted")
            shutil.copy2(tmp_path, fallback)
            print(f"Original locked; saved: {fallback}")
            return
    finally:
        tmp_path.unlink(missing_ok=True)


if __name__ == "__main__":
    build_docx()
    print(f"Saved: {DOCX_PATH}")
